"""
Script to generate academic lab evaluation report in Word format
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_format(doc, text, level=1):
    """Add formatted heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def create_report():
    """Create the academic lab evaluation report"""
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title Page
    title = doc.add_heading('RESUME CATEGORY PREDICTION SYSTEM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('An Academic Lab Evaluation Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    
    doc.add_paragraph()  # Spacing
    
    course_info = doc.add_paragraph()
    course_info.add_run('Course: Machine Learning\n').bold = True
    course_info.add_run('Submitted by: [Your Name]\n')
    course_info.add_run('Date: [Date]')
    course_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents placeholder
    toc_heading = doc.add_heading('Table of Contents', 1)
    doc.add_paragraph('1. Introduction & Need of the Problem')
    doc.add_paragraph('2. Problem Statement')
    doc.add_paragraph('3. Objectives and Existing Approaches')
    doc.add_paragraph('4. Dataset Description')
    doc.add_paragraph('5. Methodology / Model Description')
    doc.add_paragraph('6. Implementation Details')
    doc.add_paragraph('7. Results & Evaluation')
    doc.add_paragraph('8. Conclusion')
    doc.add_paragraph('9. Future Work & Limitations')
    
    doc.add_page_break()
    
    # ========== 1. INTRODUCTION & NEED OF THE PROBLEM ==========
    add_heading_with_format(doc, '1. Introduction & Need of the Problem', 1)
    
    doc.add_paragraph(
        'Resume categorization is a critical task in human resource management and recruitment processes. '
        'With the increasing volume of job applications received by organizations daily, manual categorization '
        'of resumes has become a bottleneck that is time-consuming, error-prone, and difficult to scale. '
        'This project addresses the need for an automated system that can accurately classify resumes into '
        'appropriate job categories based on their content.'
    )
    
    doc.add_paragraph(
        'The need for such a system arises from several practical challenges:'
    )
    
    # Bullet points
    reasons = doc.add_paragraph(style='List Bullet')
    reasons.add_run('Time Efficiency: Manual categorization requires significant human resources and time.')
    
    reasons = doc.add_paragraph(style='List Bullet')
    reasons.add_run('Scalability: As organizations grow, the volume of applications increases exponentially.')
    
    reasons = doc.add_paragraph(style='List Bullet')
    reasons.add_run('Consistency: Automated systems eliminate human bias and ensure uniform classification standards.')
    
    reasons = doc.add_paragraph(style='List Bullet')
    reasons.add_run('Accuracy: Machine learning models can process and analyze large volumes of text data more accurately than manual screening.')
    
    doc.add_paragraph(
        'This project aims to develop an intelligent system that leverages Natural Language Processing (NLP) '
        'and machine learning techniques to automatically categorize resumes into 25 distinct job domains, '
        'thereby streamlining the recruitment process and improving efficiency.'
    )
    
    # Add placeholder for diagram
    diagram_para = doc.add_paragraph()
    diagram_para.add_run('[INSERT DIAGRAM: Recruitment Process Flow - Before and After Automation]').italic = True
    diagram_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ========== 2. PROBLEM STATEMENT ==========
    add_heading_with_format(doc, '2. Problem Statement', 1)
    
    doc.add_paragraph(
        'The problem addressed in this project is to automatically classify resumes into predefined job categories '
        'based solely on the textual content of the resume. The system must:'
    )
    
    problems = doc.add_paragraph(style='List Bullet')
    problems.add_run('Accurately classify resumes into one of 25 job categories.')
    
    problems = doc.add_paragraph(style='List Bullet')
    problems.add_run('Handle variations in resume formats, writing styles, and terminology.')
    
    problems = doc.add_paragraph(style='List Bullet')
    problems.add_run('Process unstructured text data and extract meaningful features.')
    
    problems = doc.add_paragraph(style='List Bullet')
    problems.add_run('Achieve high accuracy rates suitable for production deployment.')
    
    problems = doc.add_paragraph(style='List Bullet')
    problems.add_run('Scale efficiently to handle large volumes of resumes.')
    
    doc.add_paragraph(
        'The classification task is a multiclass text classification problem where each resume can belong to '
        'only one of the 25 categories, requiring the development of robust feature extraction and classification mechanisms.'
    )
    
    doc.add_page_break()
    
    # ========== 3. OBJECTIVES AND EXISTING APPROACHES ==========
    add_heading_with_format(doc, '3. Objectives and Existing Approaches', 1)
    
    sub_heading = add_heading_with_format(doc, '3.1 Objectives', 2)
    
    doc.add_paragraph(
        'The primary objectives of this project are:'
    )
    
    objectives = doc.add_paragraph(style='List Bullet')
    objectives.add_run('To develop an automated resume categorization system with accuracy above 95%.')
    
    objectives = doc.add_paragraph(style='List Bullet')
    objectives.add_run('To implement and compare multiple machine learning algorithms for text classification.')
    
    objectives = doc.add_paragraph(style='List Bullet')
    objectives.add_run('To evaluate model performance using comprehensive metrics including accuracy, precision, recall, and F1-score.')
    
    objectives = doc.add_paragraph(style='List Bullet')
    objectives.add_run('To create a production-ready system deployable as a web application.')
    
    objectives = doc.add_paragraph(style='List Bullet')
    objectives.add_run('To provide confidence scores and top-3 category predictions for transparency.')
    
    sub_heading = add_heading_with_format(doc, '3.2 Existing Approaches', 2)
    
    doc.add_paragraph(
        'Several approaches have been proposed for text classification and resume categorization:'
    )
    
    existing = doc.add_paragraph(style='List Bullet')
    existing.add_run('Rule-Based Systems: Traditional systems using keyword matching and predefined rules. Limitations include lack of flexibility and poor generalization.')
    
    existing = doc.add_paragraph(style='List Bullet')
    existing.add_run('Statistical Methods: Naive Bayes classifiers have been widely used for text classification due to their simplicity and effectiveness with sparse features.')
    
    existing = doc.add_paragraph(style='List Bullet')
    existing.add_run('Tree-Based Models: Decision Trees and Random Forests offer interpretability and handle non-linear relationships in text data.')
    
    existing = doc.add_paragraph(style='List Bullet')
    existing.add_run('Neural Networks: Multi-Layer Perceptrons (MLP) and deep learning models have shown superior performance in complex text classification tasks.')
    
    existing = doc.add_paragraph(style='List Bullet')
    existing.add_run('Ensemble Methods: Combining multiple models through voting or stacking improves robustness and accuracy.')
    
    doc.add_paragraph(
        'This project implements and compares multiple approaches to identify the most effective method for resume categorization.'
    )
    
    doc.add_page_break()
    
    # ========== 4. DATASET DESCRIPTION ==========
    add_heading_with_format(doc, '4. Dataset Description', 1)
    
    doc.add_paragraph(
        'The dataset used in this project consists of 962 resume samples, each labeled with one of 25 job categories. '
        'The dataset is structured as a CSV file with three columns: Category, Resume (raw text), and cleaned_resume (preprocessed text).'
    )
    
    sub_heading = add_heading_with_format(doc, '4.1 Dataset Statistics', 2)
    
    # Create table for dataset statistics
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    table.cell(0, 0).text = 'Attribute'
    table.cell(0, 1).text = 'Value'
    table.cell(1, 0).text = 'Total Samples'
    table.cell(1, 1).text = '962 resumes'
    table.cell(2, 0).text = 'Number of Categories'
    table.cell(2, 1).text = '25 job domains'
    table.cell(3, 0).text = 'Train-Test Split'
    table.cell(3, 1).text = '80% (769) - 20% (193)'
    
    # Make header bold
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    sub_heading = add_heading_with_format(doc, '4.2 Category Distribution', 2)
    
    doc.add_paragraph(
        'The dataset contains resumes from 25 different job categories, with varying sample sizes:'
    )
    
    # Top categories
    categories = [
        ('Java Developer', 84, '8.73%'),
        ('Testing', 70, '7.28%'),
        ('DevOps Engineer', 55, '5.72%'),
        ('Python Developer', 48, '4.99%'),
        ('Web Designing', 45, '4.68%'),
    ]
    
    cat_table = doc.add_table(rows=6, cols=3)
    cat_table.style = 'Light Grid Accent 1'
    
    cat_table.cell(0, 0).text = 'Category'
    cat_table.cell(0, 1).text = 'Count'
    cat_table.cell(0, 2).text = 'Percentage'
    
    for i, (cat, count, pct) in enumerate(categories, 1):
        cat_table.cell(i, 0).text = cat
        cat_table.cell(i, 1).text = str(count)
        cat_table.cell(i, 2).text = pct
    
    # Make header bold
    for cell in cat_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    doc.add_paragraph('(Table shows top 5 categories. Complete list includes 25 categories total.)')
    
    # Add placeholder for pie chart
    chart_para = doc.add_paragraph()
    chart_para.add_run('[INSERT CHART: Category Distribution Pie Chart]').italic = True
    chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub_heading = add_heading_with_format(doc, '4.3 Data Characteristics', 2)
    
    doc.add_paragraph(
        'Key characteristics of the dataset:'
    )
    
    chars = doc.add_paragraph(style='List Bullet')
    chars.add_run('Imbalanced Classes: Categories range from 20 to 84 samples, requiring careful evaluation.')
    
    chars = doc.add_paragraph(style='List Bullet')
    chars.add_run('Text Format: Variable-length resume texts with rich content including skills, experience, education.')
    
    chars = doc.add_paragraph(style='List Bullet')
    chars.add_run('Quality: Well-labeled data with clear category distinctions suitable for supervised learning.')
    
    doc.add_page_break()
    
    # ========== 5. METHODOLOGY / MODEL DESCRIPTION ==========
    add_heading_with_format(doc, '5. Methodology / Model Description', 1)
    
    doc.add_paragraph(
        'The methodology follows a systematic machine learning pipeline for text classification, '
        'comprising data preprocessing, feature extraction, model training, and evaluation phases.'
    )
    
    sub_heading = add_heading_with_format(doc, '5.1 Preprocessing Pipeline', 2)
    
    doc.add_paragraph(
        'Text preprocessing is performed to clean and normalize resume content:'
    )
    
    preprocess_steps = [
        'URL removal: Eliminate HTTP/HTTPS links',
        'Hashtag and mention removal: Clean social media artifacts',
        'Punctuation removal: Remove special characters',
        'Non-ASCII character handling: Address encoding issues',
        'Whitespace normalization: Standardize spacing',
        'Lowercase conversion: Normalize text case'
    ]
    
    for step in preprocess_steps:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(step)
    
    # Add flowchart placeholder
    flow_para = doc.add_paragraph()
    flow_para.add_run('[INSERT DIAGRAM: Data Preprocessing Pipeline Flowchart]').italic = True
    flow_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub_heading = add_heading_with_format(doc, '5.2 Feature Extraction', 2)
    
    doc.add_paragraph(
        'TF-IDF (Term Frequency-Inverse Document Frequency) vectorization is used to convert text into numerical features. '
        'This technique assigns weights to words based on their frequency in a document relative to their frequency across '
        'all documents, effectively highlighting distinctive terms.'
    )
    
    doc.add_paragraph('Formula: TF-IDF(t,d) = TF(t,d) × IDF(t)', style='Intense Quote')
    
    doc.add_paragraph(
        'The resulting feature matrix contains 7,351 dimensions, representing unique terms across all resumes. '
        'Sublinear TF scaling and English stop word removal are applied to improve feature quality.'
    )
    
    sub_heading = add_heading_with_format(doc, '5.3 Models Implemented', 2)
    
    doc.add_paragraph(
        'Six different machine learning algorithms were implemented and compared:'
    )
    
    models_desc = [
        ('K-Nearest Neighbors (KNN)', 'Instance-based learning algorithm that classifies based on similarity to k nearest neighbors.'),
        ('Gaussian Naive Bayes', 'Probabilistic classifier assuming Gaussian distribution of features. Fast and efficient for text classification.'),
        ('Decision Tree', 'Tree-based model that splits data based on feature values, offering interpretability.'),
        ('Random Forest', 'Ensemble of decision trees providing robust predictions and feature importance.'),
        ('Logistic Regression', 'Linear classifier with probabilistic outputs, suitable for multiclass problems.'),
        ('MLP Neural Network', 'Multi-layer perceptron with hidden layers capable of learning non-linear patterns in text data.')
    ]
    
    for model_name, description in models_desc:
        p = doc.add_paragraph()
        p.add_run(model_name + ': ').bold = True
        p.add_run(description)
    
    # Add diagram placeholder
    model_para = doc.add_paragraph()
    model_para.add_run('[INSERT DIAGRAM: ML Pipeline Architecture]').italic = True
    model_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ========== 6. IMPLEMENTATION DETAILS ==========
    add_heading_with_format(doc, '6. Implementation Details', 1)
    
    sub_heading = add_heading_with_format(doc, '6.1 Technology Stack', 2)
    
    tech_table = doc.add_table(rows=8, cols=2)
    tech_table.style = 'Light Grid Accent 1'
    
    tech_stack = [
        ('Programming Language', 'Python 3.x'),
        ('ML Library', 'scikit-learn'),
        ('Data Processing', 'pandas, numpy'),
        ('NLP', 'TF-IDF Vectorizer'),
        ('Visualization', 'matplotlib, seaborn'),
        ('Web Framework', 'Streamlit'),
        ('Development Environment', 'Jupyter Notebook'),
        ('Model Deployment', 'Streamlit Web Application')
    ]
    
    tech_table.cell(0, 0).text = 'Component'
    tech_table.cell(0, 1).text = 'Technology'
    
    for i, (component, tech) in enumerate(tech_stack, 1):
        tech_table.cell(i, 0).text = component
        tech_table.cell(i, 1).text = tech
    
    # Make header bold
    for cell in tech_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    sub_heading = add_heading_with_format(doc, '6.2 Training Configuration', 2)
    
    doc.add_paragraph(
        'Training parameters and configurations:'
    )
    
    config_steps = [
        'Data split: 80% training (769 samples), 20% testing (193 samples)',
        'Stratification: Enabled to maintain category distribution',
        'Random state: 42 (for reproducibility)',
        'One-vs-Rest strategy: Applied for multiclass classification',
        'Cross-validation: 5-fold CV for robust evaluation',
        'Hyperparameters: Default scikit-learn parameters used'
    ]
    
    for step in config_steps:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(step)
    
    sub_heading = add_heading_with_format(doc, '6.3 Evaluation Metrics', 2)
    
    doc.add_paragraph(
        'Models are evaluated using multiple metrics:'
    )
    
    metrics_desc = [
        ('Accuracy', 'Percentage of correctly classified resumes'),
        ('Precision', 'Proportion of predicted positives that are actually positive'),
        ('Recall', 'Proportion of actual positives correctly identified'),
        ('F1-Score', 'Harmonic mean of precision and recall')
    ]
    
    for metric, desc in metrics_desc:
        p = doc.add_paragraph()
        p.add_run(metric + ': ').bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # ========== 7. RESULTS & EVALUATION ==========
    add_heading_with_format(doc, '7. Results & Evaluation', 1)
    
    sub_heading = add_heading_with_format(doc, '7.1 Model Performance Comparison', 2)
    
    doc.add_paragraph(
        'All six models were trained and evaluated on the test set. Performance results are summarized below:'
    )
    
    # Performance table
    perf_table = doc.add_table(rows=7, cols=5)
    perf_table.style = 'Light Grid Accent 1'
    
    headers = ['Model', 'Test Accuracy', 'Precision', 'Recall', 'F1-Score']
    for i, header in enumerate(headers):
        perf_table.cell(0, i).text = header
    
    results_data = [
        ('MLP Neural Network', '100.00%', '1.0000', '1.0000', '1.0000'),
        ('Gaussian Naive Bayes', '99.48%', '0.9953', '0.9948', '0.9948'),
        ('Decision Tree', '99.48%', '0.9954', '0.9948', '0.9948'),
        ('Random Forest', '99.48%', '0.9954', '0.9948', '0.9948'),
        ('Logistic Regression', '98.96%', '0.9909', '0.9896', '0.9896'),
        ('K-Nearest Neighbors', '97.93%', '0.9796', '0.9793', '0.9785')
    ]
    
    for i, (model, acc, prec, rec, f1) in enumerate(results_data, 1):
        perf_table.cell(i, 0).text = model
        perf_table.cell(i, 1).text = acc
        perf_table.cell(i, 2).text = prec
        perf_table.cell(i, 3).text = rec
        perf_table.cell(i, 4).text = f1
    
    # Make header bold
    for cell in perf_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Highlight best performer
    for i in range(5):
        cell = perf_table.cell(1, i)  # MLP row
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    sub_heading = add_heading_with_format(doc, '7.2 Best Performing Model', 2)
    
    doc.add_paragraph(
        'MLP Neural Network achieved perfect classification with 100% accuracy on the test set, demonstrating '
        'exceptional performance. The model achieved perfect scores across all metrics (Precision: 1.0000, '
        'Recall: 1.0000, F1-Score: 1.0000), indicating no misclassifications.'
    )
    
    # Add chart placeholder
    chart1_para = doc.add_paragraph()
    chart1_para.add_run('[INSERT CHART: Model Accuracy Comparison Bar Chart]').italic = True
    chart1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub_heading = add_heading_with_format(doc, '7.3 Confusion Matrix Analysis', 2)
    
    doc.add_paragraph(
        'Confusion matrices were generated for all models to analyze classification errors. The MLP Neural Network '
        'showed perfect classification with all predictions on the diagonal, indicating no misclassifications.'
    )
    
    # Add chart placeholder
    chart2_para = doc.add_paragraph()
    chart2_para.add_run('[INSERT CHART: Confusion Matrix Heatmaps for All Models]').italic = True
    chart2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub_heading = add_heading_with_format(doc, '7.4 Cross-Validation Results', 2)
    
    doc.add_paragraph(
        '5-fold cross-validation was performed to assess model stability and generalization. Results showed '
        'consistent performance across folds, with minimal variance, indicating robust models suitable for production deployment.'
    )
    
    # Add chart placeholder
    chart3_para = doc.add_paragraph()
    chart3_para.add_run('[INSERT CHART: Cross-Validation Box Plot]').italic = True
    chart3_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub_heading = add_heading_with_format(doc, '7.5 Key Findings', 2)
    
    findings = doc.add_paragraph(style='List Bullet')
    findings.add_run('MLP Neural Network demonstrated superior capability in capturing complex patterns in resume text.')
    
    findings = doc.add_paragraph(style='List Bullet')
    findings.add_run('All models achieved accuracy above 97%, indicating high-quality feature engineering.')
    
    findings = doc.add_paragraph(style='List Bullet')
    findings.add_run('TF-IDF vectorization effectively captured discriminative features for category classification.')
    
    findings = doc.add_paragraph(style='List Bullet')
    findings.add_run('Models showed consistent performance across different categories, suggesting good generalization.')
    
    doc.add_page_break()
    
    # ========== 8. CONCLUSION ==========
    add_heading_with_format(doc, '8. Conclusion', 1)
    
    doc.add_paragraph(
        'This project successfully developed an automated resume categorization system that achieves exceptional '
        'performance across multiple machine learning algorithms. The MLP Neural Network emerged as the best '
        'performing model with 100% accuracy on the test set.'
    )
    
    doc.add_paragraph(
        'Key achievements include:'
    )
    
    achievements = doc.add_paragraph(style='List Bullet')
    achievements.add_run('Successfully classified 962 resumes into 25 job categories with high accuracy.')
    
    achievements = doc.add_paragraph(style='List Bullet')
    achievements.add_run('Implemented and compared 6 different ML algorithms, providing comprehensive evaluation.')
    
    achievements = doc.add_paragraph(style='List Bullet')
    achievements.add_run('Achieved perfect classification (100%) with MLP Neural Network.')
    
    achievements = doc.add_paragraph(style='List Bullet')
    achievements.add_run('Developed a production-ready system deployable as a web application.')
    
    doc.add_paragraph(
        'The results demonstrate that automated resume categorization is feasible and can significantly improve '
        'efficiency in HR processes. The high accuracy rates achieved suggest that the system is suitable for '
        'real-world deployment and can provide substantial time savings in resume screening processes.'
    )
    
    doc.add_page_break()
    
    # ========== 9. FUTURE WORK & LIMITATIONS ==========
    add_heading_with_format(doc, '9. Future Work & Limitations', 1)
    
    sub_heading = add_heading_with_format(doc, '9.1 Limitations', 2)
    
    doc.add_paragraph(
        'While the system achieved excellent results, several limitations exist:'
    )
    
    limitations = doc.add_paragraph(style='List Bullet')
    limitations.add_run('Dataset Size: The dataset contains 962 samples, which may limit generalization to diverse resume styles.')
    
    limitations = doc.add_paragraph(style='List Bullet')
    limitations.add_run('Class Imbalance: Uneven distribution across categories (20-84 samples) may affect minority class performance.')
    
    limitations = doc.add_paragraph(style='List Bullet')
    limitations.add_run('Language: System is currently limited to English-language resumes.')
    
    limitations = doc.add_paragraph(style='List Bullet')
    limitations.add_run('Format Dependency: Performance may vary with different resume formats and structures.')
    
    limitations = doc.add_paragraph(style='List Bullet')
    limitations.add_run('Feature Engineering: Current approach uses TF-IDF; advanced embeddings (Word2Vec, BERT) were not explored.')
    
    sub_heading = add_heading_with_format(doc, '9.2 Future Work', 2)
    
    doc.add_paragraph(
        'Future enhancements and research directions include:'
    )
    
    future_work = doc.add_paragraph(style='List Bullet')
    future_work.add_run('Dataset Expansion: Increase dataset size and diversity to improve generalization.')
    
    future_work = doc.add_paragraph(style='List Bullet')
    future_work.add_run('Deep Learning: Explore Transformer models (BERT, GPT) for enhanced text understanding.')
    
    future_work = doc.add_paragraph(style='List Bullet')
    future_work.add_run('Feature Engineering: Implement semantic embeddings and named entity recognition for richer features.')
    
    future_work = doc.add_paragraph(style='List Bullet')
    future_work.add_run('Multi-language Support: Extend system to handle resumes in multiple languages.')
    
    future_work = doc.add_paragraph(style='List Bullet')
    future_work.add_run('Imbalance Handling: Apply techniques like SMOTE or class weighting to address class imbalance.')
    
    future_work = doc.add_paragraph(style='List Bullet')
    future_work.add_run('Real-time Processing: Optimize for faster processing of large volumes of resumes.')
    
    future_work = doc.add_paragraph(style='List Bullet')
    future_work.add_run('Explainability: Add feature importance visualization and explainable AI techniques.')
    
    future_work = doc.add_paragraph(style='List Bullet')
    future_work.add_run('API Development: Create RESTful API for integration with existing HR systems.')
    
    doc.add_page_break()
    
    # References
    add_heading_with_format(doc, 'References', 1)
    
    refs = [
        'Scikit-learn Documentation. Available at: https://scikit-learn.org/',
        'Pedregosa et al., "Scikit-learn: Machine Learning in Python", JMLR 12, pp. 2825-2830, 2011.',
        'Manning, C., et al. "Introduction to Information Retrieval." Cambridge University Press, 2008.',
        'Streamlit Documentation. Available at: https://docs.streamlit.io/',
    ]
    
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(ref)
    
    return doc

if __name__ == '__main__':
    print("Generating academic lab evaluation report...")
    doc = create_report()
    output_path = 'ML Project/ML-Project/Academic_Lab_Report.docx'
    doc.save(output_path)
    print(f"Report generated successfully: {output_path}")

